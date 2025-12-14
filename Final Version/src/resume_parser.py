"""
Resume parsing utilities for PDF, DOCX, and TXT files.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        content = "\n".join(text_parts)
        logger.debug(f"Extracted {len(content)} chars from PDF: {file_path}")
        return content.strip()
        
    except Exception as e:
        logger.error(f"Failed to parse PDF {file_path}: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        content = "\n".join(text_parts)
        logger.debug(f"Extracted {len(content)} chars from DOCX: {file_path}")
        return content.strip()
        
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise


def parse_txt(file_path: str) -> str:
    """
    Read text from TXT file.
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        File content
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        logger.debug(f"Read {len(content)} chars from TXT: {file_path}")
        return content.strip()
        
    except Exception as e:
        logger.error(f"Failed to read TXT {file_path}: {e}")
        raise


def parse_resume(file_path: str) -> Tuple[str, str]:
    """
    Parse resume file and extract text content.
    
    Supports PDF, DOCX, and TXT formats.
    
    Args:
        file_path: Path to resume file
        
    Returns:
        Tuple of (extracted_text, file_type)
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == '.pdf':
        return parse_pdf(file_path), 'pdf'
    elif extension == '.docx':
        return parse_docx(file_path), 'docx'
    elif extension in ['.txt', '.text']:
        return parse_txt(file_path), 'txt'
    else:
        # Try as plain text
        logger.warning(f"Unknown extension {extension}, trying as text: {file_path}")
        return parse_txt(file_path), 'unknown'


def get_resume_files(directory: str, recursive: bool = True) -> list:
    """
    Get all resume files from a directory.
    
    Args:
        directory: Directory to search
        recursive: Search subdirectories
        
    Returns:
        List of file paths
    """
    extensions = {'.pdf', '.docx', '.txt', '.text'}
    path = Path(directory)
    
    if not path.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")
    
    if recursive:
        files = [
            str(f) for f in path.rglob('*')
            if f.is_file() and f.suffix.lower() in extensions
        ]
    else:
        files = [
            str(f) for f in path.iterdir()
            if f.is_file() and f.suffix.lower() in extensions
        ]
    
    logger.info(f"Found {len(files)} resume files in {directory}")
    return sorted(files)


def extract_candidate_name(content: str, file_path: str) -> str:
    """
    Attempt to extract candidate name from resume content.
    Falls back to filename if extraction fails.
    
    Args:
        content: Resume text content
        file_path: Path to resume file
        
    Returns:
        Candidate name
    """
    # Try to get name from first non-empty line (common pattern)
    lines = [l.strip() for l in content.split('\n') if l.strip()]
    
    if lines:
        first_line = lines[0]
        # If first line looks like a name (2-4 words, no special chars)
        words = first_line.split()
        if 2 <= len(words) <= 4:
            if all(word.replace('-', '').replace("'", '').isalpha() for word in words):
                return first_line
    
    # Fall back to filename
    filename = Path(file_path).stem
    # Clean up filename (remove underscores, dashes)
    name = filename.replace('_', ' ').replace('-', ' ')
    return name.title()
