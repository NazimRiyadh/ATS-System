import requests
import os
from docx import Document

BASE_URL = "http://127.0.0.1:8000"

def create_docx(filename):
    try:
        doc = Document()
        doc.add_heading('John Docx', 0)
        doc.add_paragraph('Senior Java Developer')
        doc.add_paragraph('Experience: 5 years in Spring Boot.')
        doc.add_paragraph('Skills: Java, Spring, SQL.')
        doc.save(filename)
        print(f"Created {filename}")
    except Exception as e:
        print(f"Failed to create DOCX: {e}")

def create_pdf(filename):
    # Valid minimal PDF structure containing "John PDF Candidate"
    # Generated content for a simple PDF 1.1 file
    content = (
        b"%PDF-1.1\n"
        b"% \n"
        b"1 0 obj\n<<\n/Type /Catalog\n/Pages 2 0 R\n>>\nendobj\n"
        b"2 0 obj\n<<\n/Type /Pages\n/Kids [3 0 R]\n/Count 1\n/MediaBox [0 0 300 144]\n>>\nendobj\n"
        b"3 0 obj\n<<\n/Type /Page\n/Parent 2 0 R\n/Resources <<\n/Font <<\n/F1 4 0 R\n>>\n>>\n/Contents 5 0 R\n>>\nendobj\n"
        b"4 0 obj\n<<\n/Type /Font\n/Subtype /Type1\n/BaseFont /Times-Roman\n>>\nendobj\n"
        b"5 0 obj\n<<\n/Length 55\n>>\nstream\n"
        b"BT\n/F1 18 Tf\n0 0 Td\n(John PDF Candidate) Tj\nET\n"
        b"endstream\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000018 00000 n \n0000000077 00000 n \n0000000178 00000 n \n0000000302 00000 n \n0000000388 00000 n \n"
        b"trailer\n<<\n/Root 1 0 R\n/Size 6\n>>\nstartxref\n493\n%%EOF"
    )
    
    try:
        with open(filename, 'wb') as f:
            f.write(content)
        print(f"Created {filename}")
    except Exception as e:
        print(f"Failed to create PDF: {e}")

def ingest_file(filename):
    print(f"Ingesting {filename}...")
    if not os.path.exists(filename):
        print(f"File not found: {filename}")
        return

    try:
        with open(filename, 'rb') as f:
            # Files dict structure: {'file': (name, file_obj, content_type)}
            files = {'file': (filename, f, 'application/octet-stream')}
            resp = requests.post(f"{BASE_URL}/ingest", files=files)
            
            if resp.status_code == 200:
                print(f"[SUCCESS] Ingestion for {filename}: OK")
                print(resp.json())
            else:
                print(f"[FAILED] Ingestion for {filename}: {resp.status_code}")
                # Print first 200 chars of error to avoid huge log
                print(resp.text[:200])
    except Exception as e:
        print(f"[ERROR] {e}")

if __name__ == "__main__":
    create_docx("test_candidate.docx")
    create_pdf("test_candidate.pdf")
    
    print("\nStarting Ingestion Tests...")
    ingest_file("test_candidate.docx")
    ingest_file("test_candidate.pdf")
